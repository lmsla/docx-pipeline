import unittest

from docx import Document
from docx.oxml.ns import qn

from docx_pipeline.postprocess import normalize_tables


class TableGeometryTest(unittest.TestCase):
    def test_revision_table_gets_explicit_full_width_geometry(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=4)
        headers = ["修訂日期", "版號", "修訂內容", "修訂者"]
        for cell, text in zip(table.rows[0].cells, headers):
            cell.text = text
        for grid_column in list(table._tbl.tblGrid):
            table._tbl.tblGrid.remove(grid_column)
        for row in table.rows:
            for cell in row.cells:
                cell._tc.tcPr.remove(cell._tc.tcPr.tcW)

        normalize_tables(doc)

        expected_total = round(
            (doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin) / 635
        )
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        grid_widths = [int(column.get(qn("w:w"))) for column in table._tbl.tblGrid]
        expected_widths = [expected_total * weight // 100 for weight in [18, 10, 52, 20]]
        expected_widths[-1] += expected_total - sum(expected_widths)

        self.assertEqual(tbl_w.get(qn("w:type")), "dxa")
        self.assertEqual(int(tbl_w.get(qn("w:w"))), expected_total)
        tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
        self.assertEqual(tbl_ind.get(qn("w:type")), "dxa")
        self.assertEqual(int(tbl_ind.get(qn("w:w"))), 0)
        self.assertEqual(sum(grid_widths), expected_total)
        self.assertEqual(grid_widths, expected_widths)
        for row in table.rows:
            self.assertEqual(
                [int(cell._tc.tcPr.tcW.get(qn("w:w"))) for cell in row.cells],
                grid_widths,
            )

    def test_existing_non_revision_geometry_is_preserved(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "屬性"
        table.rows[0].cells[1].text = "內容"
        table_width = table._tbl.tblPr.find(qn("w:tblW"))
        table_width.set(qn("w:type"), "pct")
        table_width.set(qn("w:w"), "5000")
        for column, width in zip(table._tbl.tblGrid, [3000, 4920]):
            column.set(qn("w:w"), str(width))

        normalize_tables(doc)

        self.assertEqual(table_width.get(qn("w:type")), "pct")
        self.assertEqual(table_width.get(qn("w:w")), "5000")
        self.assertEqual(
            [int(column.get(qn("w:w"))) for column in table._tbl.tblGrid],
            [3000, 4920],
        )

    def test_decimal_zero_geometry_is_repaired(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        for column in table._tbl.tblGrid:
            column.set(qn("w:w"), "0.0")
        table._tbl.tblPr.find(qn("w:tblW")).set(qn("w:w"), "0.0")

        normalize_tables(doc)

        self.assertTrue(all(int(column.get(qn("w:w"))) > 0 for column in table._tbl.tblGrid))


if __name__ == "__main__":
    unittest.main()
