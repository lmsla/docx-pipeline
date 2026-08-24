from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory
from zipfile import ZipFile, ZIP_DEFLATED
import datetime as dt
import re
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor, Twips


A4_WIDTH = Cm(21.0)
A4_HEIGHT = Cm(29.7)
MARGIN = Inches(1.0)
CONTENT_WIDTH = Inches(6.27)
CODE_FILL = "F2F2F2"
CODE_BORDER = "D9D9D9"
EMU_PER_DXA = 635
REVISION_TABLE_HEADERS = ["修訂日期", "版號", "修訂內容", "修訂者"]
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


def strip_embedded_fonts(input_docx: Path, output_docx: Path) -> None:
    """Remove embedded font binary parts that python-docx cannot always parse.

    The Word style definitions still keep their font-family declarations; this
    only removes bundled TTF payloads and their fontTable relationships.
    """
    with ZipFile(input_docx, "r") as src, ZipFile(output_docx, "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            if item.filename.startswith("word/fonts/"):
                continue
            if item.filename == "word/_rels/fontTable.xml.rels":
                continue
            data = src.read(item.filename)
            if item.filename == "word/fontTable.xml":
                data = remove_font_embed_refs(data)
            dst.writestr(item, data)


def remove_font_embed_refs(font_table_xml: bytes) -> bytes:
    # 用字串級移除而非 ElementTree round-trip：ET 重新序列化會丟失
    # root 上的 namespace 宣告（w14/w15/w16…），但 mc:Ignorable 仍引用
    # 這些 prefix，Word 會視為格式錯誤。
    return re.sub(rb"<w:embed[A-Za-z]*\b[^>]*?/>", b"", font_table_xml)


def clear_dirty_field_flags(docx_path: Path) -> None:
    with TemporaryDirectory() as tmpdir:
        tmp_docx = Path(tmpdir) / "clean.docx"
        with ZipFile(docx_path, "r") as src, ZipFile(tmp_docx, "w", ZIP_DEFLATED) as dst:
            for item in src.infolist():
                data = src.read(item.filename)
                if item.filename.endswith(".xml"):
                    data = data.replace(b' w:dirty="true"', b"")
                dst.writestr(item, data)
        tmp_docx.replace(docx_path)


def get_or_insert_ordered(parent, tag: str, successors: tuple[str, ...]):
    """取得或建立子元素，並依 OOXML schema 順序插入（插在第一個出現的 successor 之前）。

    直接 append 會違反 schema 元素順序，Word 2019+ 開檔會跳「無法讀取的內容」修復提示。
    """
    element = parent.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        parent.insert_element_before(element, *successors)
    return element


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = get_or_insert_ordered(
        tc_pr, "w:shd",
        ("w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign", "w:hideMark"),
    )
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)


def set_table_borders(table, *, outer_color: str = "9E9E9E", inner_color: str = "D9D9D9") -> None:
    tbl_pr = table._tbl.tblPr
    borders = get_or_insert_ordered(
        tbl_pr, "w:tblBorders",
        ("w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook", "w:tblCaption", "w:tblDescription"),
    )

    for side in ("top", "left", "bottom", "right"):
        border = borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), outer_color)

    for side in ("insideH", "insideV"):
        border = borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), inner_color)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = get_or_insert_ordered(
        p_pr, "w:shd",
        ("w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr", "w:sectPr"),
    )
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = get_or_insert_ordered(
        p_pr, "w:pBdr",
        ("w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr", "w:sectPr"),
    )

    for side in ("top", "left", "bottom", "right"):
        border = p_bdr.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            p_bdr.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), color)


def set_cell_margins(cell, margin_dxa: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = get_or_insert_ordered(
        tc_pr, "w:tcMar",
        ("w:textDirection", "w:tcFitText", "w:vAlign", "w:hideMark"),
    )

    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_dxa))
        node.set(qn("w:type"), "dxa")


def set_table_cell_text(cell, text: str, *, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(12)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, 160)
    if fill:
        set_cell_shading(cell, fill)


def is_cover_metadata_table(table) -> bool:
    if len(table.rows) not in {5, 6} or len(table.columns) != 2:
        return False
    labels = [row.cells[0].text.strip() for row in table.rows]
    expected = ["文件名稱", "版本", "建立日期", "適用系統", "適用對象"]
    return labels[:5] == expected and (len(labels) == 5 or labels[5] == "撰寫者")


def table_geometry_is_usable(table) -> bool:
    column_count = max((len(row.cells) for row in table.rows), default=0)
    grid_widths = [int(column.get(qn("w:w"), "0")) for column in table._tbl.tblGrid]
    tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
    table_width = int(tbl_w.get(qn("w:w"), "0")) if tbl_w is not None else 0
    return (
        column_count > 0
        and len(grid_widths) == column_count
        and all(width > 0 for width in grid_widths)
        and table_width > 0
    )


def set_table_geometry(table, total_width: int) -> None:
    column_count = max((len(row.cells) for row in table.rows), default=0)
    if column_count == 0:
        return

    headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
    weights = [18, 10, 52, 20] if headers == REVISION_TABLE_HEADERS else [1] * column_count
    weight_total = sum(weights)
    widths = [total_width * weight // weight_total for weight in weights]
    widths[-1] += total_width - sum(widths)

    tbl_pr = table._tbl.tblPr
    tbl_w = get_or_insert_ordered(
        tbl_pr,
        "w:tblW",
        (
            "w:jc", "w:tblCellSpacing", "w:tblInd", "w:tblBorders", "w:shd",
            "w:tblLayout", "w:tblCellMar", "w:tblLook",
        ),
    )
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_width))

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_indent = get_or_insert_ordered(
        tbl_pr,
        "w:tblInd",
        ("w:tblBorders", "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook"),
    )
    table_indent.set(qn("w:type"), "dxa")
    table_indent.set(qn("w:w"), "0")

    tbl_grid = table._tbl.tblGrid
    for grid_col in list(tbl_grid):
        tbl_grid.remove(grid_col)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for index, width in enumerate(widths):
        table.columns[index].width = Twips(width)

    for row in table.rows:
        row.height = None
        for index, cell in enumerate(row.cells):
            cell.width = Twips(widths[index])
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[index]))


def set_paragraph_style_tokens(doc: Document) -> None:
    normal = None
    for style in doc.styles:
        if style.name.lower() == "normal":
            normal = style
            break
    if normal is None:
        normal = doc.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)

    tokens = {
        "Heading 1": (18, "1F3864", 18, 10),
        "Heading 2": (14, "2E74B5", 14, 8),
        "Heading 3": (12, "1F4E79", 10, 6),
        "Heading 4": (11, "333333", 8, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Source Code" in doc.styles:
        style = doc.styles["Source Code"]
        style.font.name = "Fira Mono"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Fira Mono")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Fira Mono")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(9.5)
        style.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        style.paragraph_format.left_indent = Pt(8)
        style.paragraph_format.right_indent = Pt(8)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.0


def normalize_sections(doc: Document, header_text: str) -> None:
    for section in doc.sections:
        section.page_width = A4_WIDTH
        section.page_height = A4_HEIGHT
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)
        section.different_first_page_header_footer = True
        for paragraph in section.first_page_header.paragraphs:
            paragraph.text = ""
        for paragraph in section.first_page_footer.paragraphs:
            paragraph.text = ""

        header = section.header
        if header.paragraphs:
            paragraph = header.paragraphs[0]
        else:
            paragraph = header.add_paragraph()
        paragraph.text = header_text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in paragraph.runs:
            run.font.name = "Microsoft JhengHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = ""
        paragraph.add_run("- ")
        add_page_field(paragraph)
        paragraph.add_run(" -")
        for run in paragraph.runs:
            run.font.name = "Microsoft JhengHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
            run.font.size = Pt(9)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def normalize_tables(doc: Document) -> None:
    section = doc.sections[0]
    available_width = section.page_width - section.left_margin - section.right_margin
    fallback_width = int(round(CONTENT_WIDTH / EMU_PER_DXA))
    total_width = int(round(available_width / EMU_PER_DXA)) if available_width > 0 else fallback_width

    for table in doc.tables:
        enhance_table = not is_cover_metadata_table(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        if enhance_table:
            headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
            if headers == REVISION_TABLE_HEADERS or not table_geometry_is_usable(table):
                set_table_geometry(table, total_width)
            set_table_borders(table)
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell, 180 if enhance_table else 120)
                if row_idx == 0:
                    set_cell_shading(cell, "D9EAF7")
                elif enhance_table and row_idx % 2 == 0:
                    set_cell_shading(cell, "F7F9FB")
                for paragraph in cell.paragraphs:
                    if row_idx == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = "Microsoft JhengHei"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
                        run.font.size = Pt(10)
                        if row_idx == 0:
                            run.bold = True


def normalize_images(doc: Document) -> None:
    section = doc.sections[0]
    max_width = section.page_width - section.left_margin - section.right_margin
    if max_width <= 0:
        max_width = CONTENT_WIDTH

    for shape in doc.inline_shapes:
        if shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = max_width
            shape.height = int(shape.height * ratio)

    for paragraph in doc.paragraphs:
        has_inline_image = any(
            run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict")
            for run in paragraph.runs
        )
        if not has_inline_image:
            continue
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)


def add_cover_page(doc: Document, metadata: dict[str, str]) -> None:
    title = metadata.get("project") or metadata.get("cover_title") or metadata.get("title", "")
    subtitle = metadata.get("subtitle", "")
    guide = metadata.get("guide") or metadata.get("document_type") or ""
    file_name = metadata.get("file_name") or metadata.get("title", "")
    version = metadata.get("version", "")
    created = metadata.get("date") or dt.date.today().isoformat()
    system = metadata.get("system") or metadata.get("applicable_system") or ""
    audience = metadata.get("audience") or metadata.get("applicable_audience") or ""
    author = metadata.get("author", "")

    body = doc._body._body
    original_children = list(body)
    insert_before = original_children[0] if original_children else None

    created_elements = []

    def remember_paragraph(paragraph):
        created_elements.append(paragraph._p)
        return paragraph

    p = remember_paragraph(doc.add_paragraph())
    p.paragraph_format.space_before = Pt(58)
    p.paragraph_format.space_after = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Microsoft JhengHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    r.font.size = Pt(32)
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    if subtitle:
        p = remember_paragraph(doc.add_paragraph())
        p.paragraph_format.space_after = Pt(42)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        r.bold = True
        r.font.name = "Microsoft JhengHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        r.font.size = Pt(24)
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    if guide:
        p = remember_paragraph(doc.add_paragraph())
        p.paragraph_format.space_after = Pt(95)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"- {guide} -")
        r.font.name = "Microsoft JhengHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    labels = ["文件名稱", "版本", "建立日期", "適用系統", "適用對象"]
    values = [file_name, version, created, system, audience]
    # author 為選填：缺值時不留空列，避免封面出現空白欄位
    if author:
        labels.append("撰寫者")
        values.append(author)

    table = doc.add_table(rows=len(labels), cols=2)
    created_elements.append(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row, label, value in zip(table.rows, labels, values):
        set_table_cell_text(row.cells[0], label, bold=True, fill="D9E2F3")
        set_table_cell_text(row.cells[1], value)

    p = remember_paragraph(doc.add_paragraph())
    p.paragraph_format.space_before = Pt(12)
    p.add_run().add_break(WD_BREAK.PAGE)

    for element in created_elements:
        body.remove(element)
    for element in created_elements:
        if insert_before is None:
            body.append(element)
        else:
            body.insert(body.index(insert_before), element)


def add_static_toc(doc: Document, metadata: dict[str, str]) -> None:
    title = metadata.get("title", "").strip()
    headings: list[tuple[int, str]] = []
    skipped_title = False

    for paragraph in doc.paragraphs:
        if not paragraph.style:
            continue
        style_name = paragraph.style.name
        if style_name not in {"Heading 1", "Heading 2", "Heading 3"}:
            continue
        text = paragraph.text.strip()
        if not text:
            continue
        if not skipped_title and style_name == "Heading 1" and text == title:
            skipped_title = True
            continue
        level = int(style_name.rsplit(" ", 1)[1])
        headings.append((level, text))

    if not headings:
        return

    body = doc._body._body
    insert_before = None
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name.startswith("Heading 1"):
            insert_before = paragraph._p
            break

    created_elements = []

    def remember_paragraph(paragraph):
        created_elements.append(paragraph._p)
        return paragraph

    p = remember_paragraph(doc.add_paragraph())
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("目錄")
    run.bold = True
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    for level, text in headings:
        p = remember_paragraph(doc.add_paragraph())
        p.paragraph_format.left_indent = Pt((level - 1) * 18)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = "Microsoft JhengHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        run.font.size = Pt(11 if level > 1 else 12)
        run.bold = level == 1
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64) if level == 1 else RGBColor(0x33, 0x33, 0x33)

    p = remember_paragraph(doc.add_paragraph())
    p.add_run().add_break(WD_BREAK.PAGE)

    for element in created_elements:
        body.remove(element)
    for element in created_elements:
        if insert_before is None:
            body.append(element)
        else:
            body.insert(body.index(insert_before), element)


def remove_pandoc_title_block(doc: Document, metadata: dict[str, str]) -> None:
    removable_date = metadata.get("date", "")
    body = doc._body._body
    for paragraph in list(doc.paragraphs):
        if paragraph.style and paragraph.style.name.startswith("Heading 1"):
            break
        style_name = paragraph.style.name if paragraph.style else ""
        text = paragraph.text.strip()
        if style_name in {"Title", "Subtitle"} or (removable_date and text == removable_date):
            body.remove(paragraph._p)


def normalize_code_blocks(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if not paragraph.style or paragraph.style.name != "Source Code":
            continue

        set_paragraph_shading(paragraph, CODE_FILL)
        set_paragraph_border(paragraph, CODE_BORDER)
        paragraph.paragraph_format.left_indent = Pt(8)
        paragraph.paragraph_format.right_indent = Pt(8)
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.0

        for run in paragraph.runs:
            run.font.name = "Fira Mono"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Fira Mono")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Fira Mono")
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
            run.font.size = Pt(9.5)


def add_chapter_page_breaks(doc: Document) -> None:
    first_h1_seen = False
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name == "Heading 1":
            if first_h1_seen:
                paragraph.paragraph_format.page_break_before = True
            first_h1_seen = True


def postprocess_docx(
    input_docx: Path,
    output_docx: Path,
    *,
    header_text: str,
    metadata: dict[str, str],
    add_cover: bool,
    add_toc: bool,
    clear_dirty_fields: bool,
    normalize_tables: bool,
    normalize_images: bool,
    chapter_page_breaks: bool,
) -> None:
    with TemporaryDirectory() as tmpdir:
        sanitized_input = Path(tmpdir) / "sanitized.docx"
        strip_embedded_fonts(input_docx, sanitized_input)
        doc = Document(sanitized_input)
        author = metadata.get("author", "").strip()
        if author:
            # 同步寫入 OOXML core properties，Word 檔案內容頁籤也能追到撰寫者
            doc.core_properties.author = author
            doc.core_properties.last_modified_by = author
        set_paragraph_style_tokens(doc)
        normalize_sections(doc, header_text)
        if add_cover:
            add_cover_page(doc, metadata)
        remove_pandoc_title_block(doc, metadata)
        if add_toc:
            add_static_toc(doc, metadata)
        if normalize_tables:
            normalize_tables_fn = globals()["normalize_tables"]
            normalize_tables_fn(doc)
        if normalize_images:
            normalize_images_fn = globals()["normalize_images"]
            normalize_images_fn(doc)
        normalize_code_blocks(doc)
        if chapter_page_breaks:
            add_chapter_page_breaks(doc)

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_docx)
        if clear_dirty_fields:
            clear_dirty_field_flags(output_docx)
